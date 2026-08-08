from __future__ import annotations

import hashlib
import io

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from openpyxl import Workbook
from pypdf import PdfWriter
from pypdf.generic import (
    DecodedStreamObject,
    DictionaryObject,
    NameObject,
)

from parser_sidecar.models import BlockType, BoundingBox
from parser_sidecar.parsers import (
    DOCX_MEDIA_TYPE,
    HTML_MEDIA_TYPE,
    PDF_MEDIA_TYPE,
    XLSX_MEDIA_TYPE,
    BlockDraft,
    ParsedContent,
    _feature_gate_list_text,
    _looks_like_pdf_code,
    _merge_pdf_logical_blocks,
    _order_pdf_page_candidates,
    _pdf_admonition,
    _pdf_heading_level,
    _reconstruct_pdf_text,
    parse_document,
)
from parser_sidecar.service import _normalize_document


def test_docx_preserves_heading_paths_and_table_text() -> None:
    document = Document()
    document.core_properties.title = "Contract Demo"
    document.add_heading("Overview", level=1)
    document.add_paragraph("A normalized paragraph.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Alpha"
    table.cell(1, 1).text = "42"
    buffer = io.BytesIO()
    document.save(buffer)

    parsed = parse_document(
        buffer.getvalue(),
        source_name="demo.docx",
        content_type=DOCX_MEDIA_TYPE,
        parser_profile="AUTO",
        options={},
        max_archive_bytes=32 * 1024 * 1024,
    )

    assert parsed.engine == "python-docx"
    assert parsed.title == "Contract Demo"
    assert [block.type for block in parsed.blocks] == [
        BlockType.HEADING,
        BlockType.PARAGRAPH,
        BlockType.TABLE,
    ]
    assert parsed.blocks[1].heading_path == ["Contract Demo", "Overview"]
    assert "| Name | Value |" in parsed.blocks[2].text
    assert "| Alpha | 42 |" in parsed.blocks[2].text


def test_docx_recognizes_source_code_paragraph_style() -> None:
    document = Document()
    document.styles.add_style("Source Code", WD_STYLE_TYPE.PARAGRAPH)
    document.add_heading("Examples", level=1)
    document.add_paragraph('{"lower_bound":"Alice"}', style="Source Code")
    buffer = io.BytesIO()
    document.save(buffer)

    parsed = parse_document(
        buffer.getvalue(),
        source_name="example.docx",
        content_type=DOCX_MEDIA_TYPE,
        parser_profile="AUTO",
        options={},
        max_archive_bytes=32 * 1024 * 1024,
    )

    assert parsed.blocks[1].type == BlockType.CODE
    assert parsed.blocks[1].attributes["style"] == "Source Code"


def test_markdown_separates_heading_anchor_from_display_and_path() -> None:
    parsed = parse_document(
        b"# Hive Catalog {#hive-catalog}\n\n## Metadata cache {#410-meta-cache}\n\nDetails.\n",
        source_name="hive.md",
        content_type="text/markdown",
        parser_profile="AUTO",
        options={},
        max_archive_bytes=32 * 1024 * 1024,
    )

    assert parsed.title == "Hive Catalog"
    assert parsed.blocks[0].text == "Hive Catalog"
    assert parsed.blocks[0].attributes["anchor"] == "hive-catalog"
    assert parsed.blocks[1].text == "Metadata cache"
    assert parsed.blocks[1].heading_path == ["Hive Catalog", "Metadata cache"]
    assert parsed.blocks[1].attributes["anchor"] == "410-meta-cache"


def test_xlsx_emits_sheet_heading_and_contiguous_tables() -> None:
    workbook = Workbook()
    workbook.properties.title = "Revenue Book"
    sheet = workbook.active
    sheet.title = "Revenue"
    sheet.append(["Region", "Amount"])
    sheet.append(["East", 12])
    sheet.append([])
    sheet.append(["Region", "Amount"])
    sheet.append(["West", 18])
    buffer = io.BytesIO()
    workbook.save(buffer)

    parsed = parse_document(
        buffer.getvalue(),
        source_name="revenue.xlsx",
        content_type=XLSX_MEDIA_TYPE,
        parser_profile="LIGHTWEIGHT",
        options={},
        max_archive_bytes=32 * 1024 * 1024,
    )

    assert parsed.engine == "openpyxl"
    assert [block.type for block in parsed.blocks] == [
        BlockType.HEADING,
        BlockType.TABLE,
        BlockType.TABLE,
    ]
    assert parsed.blocks[1].heading_path == ["Revenue Book", "Revenue"]
    assert parsed.blocks[1].attributes["cellRange"] == "A1:B2"
    assert parsed.blocks[2].attributes["cellRange"] == "A4:B5"
    assert "| West | 18 |" in parsed.blocks[2].text


def test_layout_pdf_extracts_text_with_page_number_and_coordinates() -> None:
    content = _simple_pdf("Quarterly Report")

    parsed = parse_document(
        content,
        source_name="quarterly.pdf",
        content_type=PDF_MEDIA_TYPE,
        parser_profile="AUTO",
        options={},
        max_archive_bytes=32 * 1024 * 1024,
    )

    assert parsed.engine == "pymupdf"
    assert parsed.metadata["pageCount"] == 1
    assert parsed.blocks[0].page_number == 1
    assert parsed.blocks[0].bounding_box is not None
    assert "Quarterly Report" in parsed.blocks[0].text


def test_pdf_layout_orders_two_columns_without_moving_full_width_heading() -> None:
    candidates = [
        _pdf_candidate("heading", (40.0, 20.0, 560.0, 50.0), 0),
        _pdf_candidate("left-1", (40.0, 90.0, 270.0, 130.0), 1),
        _pdf_candidate("right-1", (330.0, 95.0, 560.0, 135.0), 2),
        _pdf_candidate("left-2", (40.0, 150.0, 270.0, 190.0), 3),
        _pdf_candidate("right-2", (330.0, 155.0, 560.0, 195.0), 4),
    ]

    ordered = _order_pdf_page_candidates(candidates, page_width=600.0)

    assert [candidate["text"] for candidate in ordered] == [
        "heading",
        "left-1",
        "left-2",
        "right-1",
        "right-2",
    ]
    assert all(candidate["readingOrder"] == "TWO_COLUMN" for candidate in ordered)
    assert [candidate["readingOrderIndex"] for candidate in ordered] == list(range(5))


def test_pdf_layout_keeps_single_column_indentation_in_visual_order() -> None:
    candidates = [
        _pdf_candidate("paragraph", (40.0, 40.0, 560.0, 80.0), 0),
        _pdf_candidate("indented-list", (90.0, 95.0, 520.0, 135.0), 1),
        _pdf_candidate("paragraph-2", (40.0, 150.0, 560.0, 190.0), 2),
    ]

    ordered = _order_pdf_page_candidates(candidates, page_width=600.0)

    assert [candidate["text"] for candidate in ordered] == [
        "paragraph",
        "indented-list",
        "paragraph-2",
    ]


def test_pdf_structure_rules_join_visual_wraps_without_flattening_code() -> None:
    assert (
        _reconstruct_pdf_text(["推荐使", "用 Python UDF。"], preserve_lines=False)
        == "推荐使用 Python UDF。"
    )
    assert (
        _reconstruct_pdf_text(
            ["CREATE FUNCTION demo()", "RETURNS INT", "AS $$"], preserve_lines=True
        )
        == "CREATE FUNCTION demo()\nRETURNS INT\nAS $$"
    )
    assert _looks_like_pdf_code(
        ["CREATE FUNCTION demo()", "RETURNS INT", "AS $$"], [11.0, 11.0, 11.0]
    )


def _pdf_candidate(
    text: str,
    bbox: tuple[float, float, float, float],
    index: int,
) -> dict[str, object]:
    return {
        "text": text,
        "bbox": bbox,
        "pageBlockIndex": index,
    }


def test_pdf_code_detection_requires_density_for_multiline_text() -> None:
    assert _looks_like_pdf_code(["SELECT * FROM documents;"], [11.0])
    assert not _looks_like_pdf_code(
        [
            "消费者参数说明如下。",
            "timeout_ms = 5",
            "该参数表示轮询等待时间，设置后仍会正常返回。",
        ],
        [11.0, 11.0, 11.0],
    )


def test_pdf_code_detection_recognizes_sql_continuation_and_cli_output() -> None:
    assert _looks_like_pdf_code([") t", "WHERE enabled = 1", "ORDER BY created_at"], [11.0] * 3)
    assert _looks_like_pdf_code(
        ["mysql> SHOW ROUTINE LOAD;", "Id: 10001", "Name: example_job"],
        [11.0] * 3,
    )


def test_pdf_code_reconstruction_repairs_visual_line_wraps_only() -> None:
    repaired = _reconstruct_pdf_text(
        [
            '{"max_batch_rows":"200000","timezone":"America/New_York","send_batch_parallelism":',
            '"1","load_to_single_tablet":"false","column_separator":"\\,","line_delimiter":"\\',
            'n","current_concurrent_number":"1","partial_columns":"false","merge_type"',
            ':"APPEND","exec_mem_limit":"2147483648","strict_mode":"false","jsonpaths":"","max_b',
            'atch_interval":"10","max_batch_size":"104857600","desired_concurrent_numbe',
            'r":"5","max_error_number":"0","max_filter_ratio":"1.',
            '0"}',
        ],
        preserve_lines=True,
    )
    assert "max_b\natch_interval" not in repaired
    assert "concurrent_numbe\nr" not in repaired
    assert 'max_filter_ratio":"1.0"' in repaired
    assert (
        _reconstruct_pdf_text(
            ["SELECT value", "FROM metrics", "WHERE enabled = 1"],
            preserve_lines=True,
        )
        == "SELECT value\nFROM metrics\nWHERE enabled = 1"
    )


def test_pdf_code_blocks_do_not_merge_across_pages() -> None:
    blocks = [
        BlockDraft(
            type=BlockType.CODE,
            text="SELECT 1;",
            page_number=1,
            heading_path=["Guide"],
            attributes={"language": "sql"},
        ),
        BlockDraft(
            type=BlockType.CODE,
            text="SELECT 2;",
            page_number=2,
            heading_path=["Guide"],
            attributes={"language": "sql"},
        ),
    ]

    merged = _merge_pdf_logical_blocks(blocks)

    assert [block.text for block in merged] == ["SELECT 1;", "SELECT 2;"]


def test_pdf_code_continuation_merges_from_page_bottom_to_next_page_top() -> None:
    blocks = [
        BlockDraft(
            type=BlockType.CODE,
            text="SELECT value\nFROM metrics\nAND $__timeFilter(timestamp)",
            page_number=9,
            heading_path=["Guide", "Dashboard"],
            bounding_box=BoundingBox(x=40, y=620, width=500, height=120),
            attributes={"language": "sql", "pageHeight": 792},
        ),
        BlockDraft(
            type=BlockType.CODE,
            text=") t\nWHERE value > 0\nORDER BY time",
            page_number=10,
            heading_path=["Guide", "Dashboard"],
            bounding_box=BoundingBox(x=40, y=45, width=500, height=90),
            attributes={"language": "sql", "pageHeight": 792},
        ),
        BlockDraft(
            type=BlockType.PARAGRAPH,
            text="要替换为其他 Counter 指标，请修改查询。",
            page_number=10,
            heading_path=["Guide", "Dashboard"],
            bounding_box=BoundingBox(x=40, y=150, width=500, height=35),
            attributes={"pageHeight": 792},
        ),
    ]

    merged = _merge_pdf_logical_blocks(blocks)

    assert len(merged) == 2
    assert merged[0].type == BlockType.CODE
    assert merged[0].attributes["pageEnd"] == 10
    assert merged[0].text.endswith("ORDER BY time")
    assert merged[1].text == "要替换为其他 Counter 指标，请修改查询。"


def test_pdf_cli_output_continuation_merges_across_pages() -> None:
    blocks = [
        BlockDraft(
            type=BlockType.PARAGRAPH,
            text=(
                "mysql> SHOW ROUTINE LOAD;\nId: 10001\n"
                'JobProperties: {"max_batch_rows":"200000","merge_type"'
            ),
            page_number=15,
            heading_path=["Routine load manual"],
            bounding_box=BoundingBox(x=40, y=600, width=500, height=145),
            attributes={"pageHeight": 792},
        ),
        BlockDraft(
            type=BlockType.CODE,
            text=':"APPEND",\n"max_batch_interval":"10"',
            page_number=16,
            heading_path=["Routine load manual"],
            bounding_box=BoundingBox(x=40, y=40, width=500, height=70),
            attributes={"pageHeight": 792},
        ),
    ]

    merged = _merge_pdf_logical_blocks(blocks)

    assert len(merged) == 1
    assert merged[0].type == BlockType.CODE
    assert '"merge_type":"APPEND"' in merged[0].text


def test_pdf_structure_rules_do_not_treat_lists_or_code_as_headings() -> None:
    base = {"fontSize": 12.0, "bold": True, "boldRatio": 1.0}
    assert _pdf_heading_level({**base, "text": "1. 启动服务"}, 12.0) is None
    assert _pdf_heading_level({**base, "text": "SELECT * FROM documents;"}, 12.0) is None
    assert _pdf_heading_level({**base, "text": "部署要求"}, 12.0) == 3


def test_pdf_admonition_removes_docusaurus_markers_and_duplicate_label() -> None:
    assert _pdf_admonition(":::caution 注意 请先备份。 :::") == (
        "caution",
        "注意：请先备份。",
    )


def test_markdown_admonition_is_normalized_without_flattening_inner_code() -> None:
    parsed = parse_document(
        b"# Guide\n\n:::tip Notice\n\nUse this command:\n\n```bash\nrun-demo\n```\n\n:::\n",
        source_name="guide.md",
        content_type="text/markdown",
        parser_profile="AUTO",
        options={},
        max_archive_bytes=32 * 1024 * 1024,
    )

    assert all(":::" not in block.text for block in parsed.blocks)
    marked = [block for block in parsed.blocks if block.attributes.get("admonition") == "tip"]
    assert marked
    assert any(block.type == BlockType.CODE for block in marked)


def test_markdown_normalization_preserves_nested_list_indentation() -> None:
    source = b"""# Catalog

* `{MetaStoreProperties}`

  Metadata service configuration:

  * `hms`: standard Hive Metastore
  * `glue`: AWS Glue
  * `dlf`: Aliyun DLF
"""
    parsed = parse_document(
        source,
        source_name="catalog.md",
        content_type="text/markdown",
        parser_profile="AUTO",
        options={},
        max_archive_bytes=32 * 1024 * 1024,
    )

    normalized = _normalize_document(
        source,
        source_name="catalog.md",
        requested_content_type="text/markdown",
        parser_profile="AUTO",
        parsed=parsed,
    )

    assert "\n\n  Metadata service configuration:" in normalized.normalized_markdown
    assert "\n\n  * `hms`: standard Hive Metastore" in normalized.normalized_markdown
    assert any(block.text.startswith("  * `hms`") for block in normalized.blocks)


def test_feature_gate_entries_become_semantic_lists_without_mutating_source_text() -> None:
    source = b"""# Kubelet

Feature gates:

FeatureOne=true|false (BETA - default=false)

kube:FeatureTwo=true|false (ALPHA - default=true)
"""
    parsed = parse_document(
        source,
        source_name="kubelet.md",
        content_type="text/markdown",
        parser_profile="AUTO",
        options={},
        max_archive_bytes=32 * 1024 * 1024,
    )
    feature_blocks = [
        block
        for block in parsed.blocks
        if block.attributes.get("structureDetected") == "FEATURE_GATE_LIST"
    ]

    assert len(feature_blocks) == 2
    assert all(block.type == BlockType.LIST for block in feature_blocks)
    assert all(not block.text.startswith("- ") for block in feature_blocks)

    normalized = _normalize_document(
        source,
        source_name="kubelet.md",
        requested_content_type="text/markdown",
        parser_profile="AUTO",
        parsed=parsed,
    )

    assert "- FeatureOne=true|false (BETA - default=false)" in normalized.normalized_markdown
    assert "- kube:FeatureTwo=true|false (ALPHA - default=true)" in normalized.normalized_markdown
    assert normalized.quality.metrics["featureGateListBlocks"] == 2
    assert normalized.quality.metrics["featureGateListItems"] == 2
    assert normalized.quality.metrics["listBlockCount"] == 2


def test_feature_gate_detection_preserves_an_explanatory_prefix() -> None:
    text = (
        "You can configure these gates:\nFeatureOne=true|false (BETA) FeatureTwo=true|false (ALPHA)"
    )

    assert _feature_gate_list_text(text) == (
        "You can configure these gates:\n"
        "FeatureOne=true|false (BETA)\n"
        "FeatureTwo=true|false (ALPHA)"
    )


def test_docx_recognizes_feature_gate_paragraphs_as_lists() -> None:
    document = Document()
    document.add_heading("Kube API server", level=1)
    document.add_paragraph("FeatureOne=true|false (BETA - default=false)")
    buffer = io.BytesIO()
    document.save(buffer)

    parsed = parse_document(
        buffer.getvalue(),
        source_name="kube-apiserver.docx",
        content_type=DOCX_MEDIA_TYPE,
        parser_profile="AUTO",
        options={},
        max_archive_bytes=32 * 1024 * 1024,
    )

    assert parsed.blocks[1].type == BlockType.LIST
    assert parsed.blocks[1].text == "FeatureOne=true|false (BETA - default=false)"
    assert parsed.blocks[1].attributes["structureDetected"] == "FEATURE_GATE_LIST"
    assert parsed.blocks[1].attributes["itemCount"] == 1


def test_normalization_produces_stable_hashes_versions_and_offsets() -> None:
    source = b"# Guide\n\n## Start\n\nHello, world."
    parsed = parse_document(
        source,
        source_name="guide.md",
        content_type="text/markdown",
        parser_profile="AUTO",
        options={},
        max_archive_bytes=32 * 1024 * 1024,
    )

    first = _normalize_document(
        source,
        source_name="guide.md",
        requested_content_type="text/markdown",
        parser_profile="AUTO",
        parsed=parsed,
    )
    second = _normalize_document(
        source,
        source_name="guide.md",
        requested_content_type="text/markdown",
        parser_profile="AUTO",
        parsed=parsed,
    )

    canonical_text = "\n\n".join(block.text for block in first.blocks)
    assert first.content_hash == hashlib.sha256(source).hexdigest()
    assert [block.block_id for block in first.blocks] == [block.block_id for block in second.blocks]
    assert (
        first.metadata["normalizedTextHash"]
        == hashlib.sha256(canonical_text.encode("utf-8")).hexdigest()
    )
    for block in first.blocks:
        encoded_prefix = canonical_text.encode("utf-16-le")[: block.source_start * 2]
        encoded_value = canonical_text.encode("utf-16-le")[
            block.source_start * 2 : block.source_end * 2
        ]
        assert len(encoded_prefix) == block.source_start * 2
        assert encoded_value.decode("utf-16-le") == block.text
        assert block.source_offset_unit == "UTF16_CODE_UNIT"
        assert block.attributes["blockVersion"] == "2.0"
        assert (
            block.attributes["contentHash"]
            == hashlib.sha256(block.text.encode("utf-8")).hexdigest()
        )

    assert first.normalized_markdown.startswith("# Guide")
    assert first.quality.status.value == "PASS"
    assert first.quality.metrics["titleBlockCount"] == 1
    assert first.quality.metrics["headingBlockCount"] == 1
    assert first.quality.metrics["paragraphBlockCount"] == 1
    assert first.quality.metrics["headingPathCoverage"] == 1.0


def test_quality_gate_rejects_content_that_is_too_short_for_retrieval() -> None:
    source = b"tiny"
    parsed = parse_document(
        source,
        source_name="tiny.txt",
        content_type="text/plain",
        parser_profile="AUTO",
        options={},
        max_archive_bytes=32 * 1024 * 1024,
    )

    normalized = _normalize_document(
        source,
        source_name="tiny.txt",
        requested_content_type="text/plain",
        parser_profile="AUTO",
        parsed=parsed,
    )

    assert normalized.quality.status.value == "FAIL"
    assert normalized.quality.score == 20
    assert [issue.code for issue in normalized.quality.issues] == ["CONTENT_TOO_SHORT"]


def test_quality_gate_warns_when_pdf_contains_source_markup_leakage() -> None:
    parsed = ParsedContent(
        title="Broken conversion",
        media_type=PDF_MEDIA_TYPE,
        engine="pymupdf",
        metadata={"pageCount": 1},
        blocks=[
            BlockDraft(
                type=BlockType.CODE,
                text=(
                    "def is_prime(n):\n"
                    "    if n is None or n **注意**：\n"
                    "> - The generated PDF joined an admonition to this code block."
                ),
                page_number=1,
            )
        ],
    )

    normalized = _normalize_document(
        b"%PDF-fixture",
        source_name="broken.pdf",
        requested_content_type=PDF_MEDIA_TYPE,
        parser_profile="AUTO",
        parsed=parsed,
    )

    assert normalized.quality.status.value == "WARNING"
    assert normalized.quality.score == 75
    assert normalized.quality.metrics["sourceMarkupLeakageBlocks"] == 1
    assert [issue.code for issue in normalized.quality.issues] == ["SOURCE_MARKUP_LEAKAGE"]


def test_html_preserves_structure_and_removes_active_or_duplicate_content() -> None:
    content = """
    <!doctype html>
    <html><head><title>企业部署手册</title><style>.hidden{display:none}</style></head>
    <body>
      <nav>站点导航</nav><script>alert('bad')</script>
      <h1>企业部署手册</h1>
      <h2>安装要求</h2>
      <p>服务器至少需要八核处理器和十六 GB 内存。</p>
      <ul><li>准备数据库<ul><li>创建业务账号</li></ul></li><li>启动服务</li></ul>
      <pre><code>systemctl start example</code></pre>
      <table><tr><th>级别</th><th>响应时间</th></tr><tr><td>P1</td><td>十分钟</td></tr></table>
      <figure>
        <img src="https://example.invalid/a.png" alt="部署拓扑">
        <figcaption>生产部署拓扑</figcaption>
      </figure>
    </body></html>
    """.encode()

    parsed = parse_document(
        content,
        source_name="deployment.html",
        content_type=HTML_MEDIA_TYPE,
        parser_profile="AUTO",
        options={},
        max_archive_bytes=32 * 1024 * 1024,
    )

    assert parsed.engine == "lxml-html"
    assert parsed.title == "企业部署手册"
    assert [block.type for block in parsed.blocks] == [
        BlockType.TITLE,
        BlockType.HEADING,
        BlockType.PARAGRAPH,
        BlockType.LIST,
        BlockType.LIST,
        BlockType.LIST,
        BlockType.CODE,
        BlockType.TABLE,
        BlockType.CAPTION,
        BlockType.CAPTION,
    ]
    text = "\n".join(block.text for block in parsed.blocks)
    assert "站点导航" not in text
    assert "alert" not in text
    assert text.count("创建业务账号") == 1
    assert "| P1 | 十分钟 |" in text
    assert parsed.blocks[2].heading_path == ["企业部署手册", "安装要求"]


def test_html_is_detected_from_extension_when_browser_omits_media_type() -> None:
    parsed = parse_document(
        b"<html><body><h1>Guide</h1><p>Content</p></body></html>",
        source_name="guide.htm",
        content_type="application/octet-stream",
        parser_profile="AUTO",
        options={},
        max_archive_bytes=32 * 1024 * 1024,
    )

    assert parsed.media_type == HTML_MEDIA_TYPE
    assert parsed.title == "Guide"


def _simple_pdf(text: str) -> bytes:
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    font_reference = writer._add_object(font)
    page[NameObject("/Resources")] = DictionaryObject(
        {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font_reference})}
    )
    stream = DecodedStreamObject()
    escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream.set_data(f"BT /F1 12 Tf 72 720 Td ({escaped}) Tj ET".encode("ascii"))
    page[NameObject("/Contents")] = writer._add_object(stream)
    output = io.BytesIO()
    writer.write(output)
    return output.getvalue()
